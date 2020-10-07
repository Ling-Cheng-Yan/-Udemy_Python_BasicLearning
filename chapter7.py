'''
大部分你想到的功能都有人寫好了，所以使用第三方套件是寫程式的關鍵。
(P.S.大部分的套件也使用其他的套件，所以寫程式有點像是堆積木和樂高，不斷地建立在現有的基石之上，不用浪費時間寫一堆早就有的功能，讓我們可以很快地開發出程式想要的功能。)
'''


'''
關於做excel的套件(google打python excel，找到openpyxl就可以去裝)
'''
from openpyxl import Workbook
wb = Workbook() #複習: wb是個物件，他的型別是Workbook()。另外，通常載入套件後，會馬上對此創造出一個物件，例如現在創造出wb。

# grab the active worksheet
ws = wb.active

# Data can be assigned directly to cells
ws['A1'] = 42
ws['B1'] = 'Allen'

# Rows can also be appended
ws.append([1, 2, 3])

# Python types will automatically be converted
import datetime
ws['A2'] = datetime.datetime.now()

# Save the file
wb.save("sample.xlsx")


'''
關於word檔套件(google打 python word，看到python-docx就是了)
'''
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

#document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')


'''
關於圖表的套件(google打 python matplotlib。以下為其中一套件示範。)
'''
#不直接創造出物件，比較特別。
import matplotlib.pyplot as plt #as就是把套件或模組簡稱，後面就不用打那麼長了。
plt.plot([1, 2, 3, 4])
plt.ylabel('some numbers')
plt.show()

plt.plot([1, 2, 3, 4], [1, 4, 9, 16], 'ro') #'ro'就是某種功能的意思，像是with open('r')，有時候function會用字串來代表功能。
plt.axis([0, 6, 0, 20])
#plt.show()
plt.savefig('123.png') #儲存圖片，如果有show就會存不起來，所以上面要先註解掉。


'''
如何寫自己的程式
【怎麼寫自己的程式】其實你們已經有所概念了，
學到這邊，相信你們對寫程式已經有一定的了解。
基礎會了，第三方套件的強大也見識到了。

其實【怎麼寫自己的程式】沒有什麼祕密SOP，
我只能盡量給大家大格局的流程模擬：

1. 想清楚自己的程式要有什麼功能
先有願景 ，input output是什麼，   先稍微思考開發的階段性，程式的發展性，想做到什麼程度
例如思考要不    要先做基本款，以後再增加網頁介面，這時候在思考的是一種【步驟】的概念
可能思考的問題：
1. 程式需要互動嗎？
   這個程式需要互動操作嗎？（例如透過CMD輸入資料）
   還是只要執行就好了 （例如執行後就自動開始下載某些資料）

2. 程式需要哪些功能？
   例如需要儲存資料嗎 (阿我可能就要想我要用哪種格式儲存，文字檔嗎還是需要使用到資料庫
   (基礎課程沒有教資料庫，但同學可以上網自學也可，先從什麼是資料庫查起) 

2. 把問題拆解 (你想達到的功能就是待解決的問題)
例如你想寫程式來定期自動貼文到臉書，
聽到這樣的問題，我會想到兩個部分：
1. 如何貼文到臉書
2. 如何定期
大家要練習這樣思維上的拆解！

3. 先把程式的想法拋到GOOGLE看看有什麼東西
說不定已經有現成的程式或是教學，完全跟你想做的功能一樣，
先google 'python facebook 貼文' 就會找到不少教學
再來就要解決如何定期貼文，google 'python 定期' 又會查到不少教學。
那你這兩個都達的到，不是組合在一起就完成了嗎！

4. 尋找相關套件
google你的想法時，出現很多現成的套件，例如跟臉書互動的套件，
這時你就該去看看這個套件的＂官網＂，先稍微看一下介紹，安裝，跟使用範例等等。
看起來很順眼，就準備好使用它了！

5. 初始能查的東西查得差不多了，可以開始動工了！
動工的第一步就是建立GitHub專案，隨著你的開發，做好辦本控管。
如果使用到一些套件，也許是邊看套件的教學，邊做。
如果覺得程式行數開始越來越多了，可以看refactor（重構程式）那些影片看看怎麼把程式碼用function收納好。
這部分比較屬於嘗試，修改，嘗試，修改。
沒有任何公認的SOP表示，寫程式一定要直接先寫function，
只要你能把你的程式做到你想要的功能，就已經任務達成了，
當你發現一件事情要重複做一百次，你自然而然會發現太麻煩，所以必須寫成function或是用到迴圈，
所以沒有任何＂標準步驟＂，都是很自然的a不行就試b, b不行就換c。
一般現實生活中，每個寫程式的人寫法習慣等等其實都相差許多，有的人喜歡用這個軟體，有的人喜歡那個，
所以越是了解這個領域，越是會發現很多很多的東西都在激辯當中，各種事情都各有優缺，很少很少事情已經有 公布說＂這就是最好的做法＂。
所以大家勇於去嘗試就對了！



最後送給大家一句，你們都已經知道的話:

所有東西都Google的到，只要有心，你絕對做得出你想做的程式!
'''